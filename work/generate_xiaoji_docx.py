from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/wangpengfei/Documents/Codex/2026-07-06/xian-b")
OUT = ROOT / "outputs" / "小迹_旅行手办_设计理念与技术方案.docx"


def set_run_font(run, name="Arial", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))


def set_paragraph_format(paragraph, *, before=0, after=6, line=1.1, align=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if align is not None:
        paragraph.alignment = align


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        edge_data = kwargs.get(edge)
        if edge_data:
            element = borders.find(qn(f"w:{edge}"))
            if element is None:
                element = OxmlElement(f"w:{edge}")
                borders.append(element)
            for key, value in edge_data.items():
                element.set(qn(f"w:{key}"), str(value))


def set_table_layout(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, header_fill=None, body_fill=None):
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            if row_idx == 0 and header_fill:
                set_cell_shading(cell, header_fill)
            elif body_fill:
                set_cell_shading(cell, body_fill)
            set_cell_border(
                cell,
                top={"val": "single", "sz": 8, "color": "D0D5DD"},
                bottom={"val": "single", "sz": 8, "color": "D0D5DD"},
                left={"val": "single", "sz": 8, "color": "D0D5DD"},
                right={"val": "single", "sz": 8, "color": "D0D5DD"},
            )


def put_text(cell, text, *, bold=False, size=10.5, color="1F2937", align=None):
    p = cell.paragraphs[0]
    if p.runs:
        for run in list(p.runs):
            run.clear()
    else:
        p = cell.paragraphs[0]
    p.alignment = align or WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12 if level == 2 else 8)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 6 if level == 2 else 4)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, bold=True, color="2E74B5")
    elif level == 2:
        set_run_font(run, size=13, bold=True, color="2E74B5")
    else:
        set_run_font(run, size=12, bold=True, color="1F4D78")
    return p


def add_body(doc, text, *, before=0, after=6, size=11, color="22231F", align=None):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    set_paragraph_format(p, before=before, after=after, line=1.1, align=align)
    r = p.add_run(text)
    set_run_font(r, size=size, color=color)
    return p


def add_callout(doc, lead, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_layout(table, [6.5])
    style_table(table, body_fill="F4F6F9")
    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    lead_run = p.add_run(lead)
    set_run_font(lead_run, size=11.5, bold=True, color="0B2545")
    body_run = p.add_run(text)
    set_run_font(body_run, size=11, color="374151")
    return table


def add_key_value_table(doc, rows, widths):
    table = doc.add_table(rows=1, cols=len(widths))
    set_table_layout(table, widths)
    style_table(table, header_fill="E8EEF5")
    hdr = table.rows[0].cells
    for idx, label in enumerate(rows[0]):
        put_text(hdr[idx], label, bold=True, size=10.5, color="0B2545", align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows[1:]:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            put_text(cells[idx], text, size=10.2, color="2B2F33", align=WD_ALIGN_PARAGRAPH.LEFT)
    return table


def set_doc_defaults(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, side, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("22231F")

    for name, size, color in (("Title", 24, "0B2545"), ("Subtitle", 11, "6B7280"), ("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5"), ("Heading 3", 12, "1F4D78")):
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style._element.rPr.rFonts.set(qn("w:cs"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    doc.core_properties.title = "小迹旅行手办产品设计理念与技术方案"
    doc.core_properties.subject = "产品设计与技术方案"
    doc.core_properties.author = "Codex"


def main():
    doc = Document()
    set_doc_defaults(doc)

    # Title block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    title = p.add_run("小迹旅行手办产品设计理念与技术方案")
    set_run_font(title, size=24, bold=True, color="0B2545")

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(14)
    p2.paragraph_format.line_spacing = 1.0
    subtitle = p2.add_run("Web 上传版 demo · 面向 MVP 落地 · 2026-07-07")
    set_run_font(subtitle, size=11, color="6B7280")

    add_callout(
        doc,
        "方案结论：",
        "本产品不是一个普通旅行相册，而是把“旅行手办”作为叙事主角，用时间线把照片、地点、餐饮和航班组织成可编辑的旅行资产，再在这份结构化数据上生成 plog、vlog 和其他表达层内容。"
    )

    add_heading(doc, "1. 产品定位", level=1)
    add_body(
        doc,
        "用户拍照或上传图片后，系统先把素材整理成事件，再把事件沉淀成故事。最终呈现的不是零散图片，而是一段可回看、可编辑、可分享、可再次生成的旅行记录。"
    )
    add_body(
        doc,
        "这意味着产品要同时解决两件事：一是降低记录成本，让用户“上传就能成稿”；二是提高内容完成度，让每一次旅行都能以手办为中心形成明确的叙事线。"
    )

    add_heading(doc, "2. 设计理念", level=1)
    add_heading(doc, "2.1 以手办为固定主角", level=2)
    add_body(
        doc,
        "传统相册记录的是照片，手办产品记录的是“我和这个手办一起完成了什么”。主角稳定后，用户会自然把不同城市、不同季节、不同时间段的旅行串成同一个角色视角，这也是产品形成记忆点的关键。"
    )
    add_heading(doc, "2.2 以时间线替代相册平铺", level=2)
    add_body(
        doc,
        "旅行天然是时间序列：出发、到达、入住、吃饭、景点、返程。将素材按时间线组织，能显著降低编辑难度，也更适合后续生成文案、卡片和短视频脚本。"
    )
    add_heading(doc, "2.3 以“上传即整理”作为第一体验", level=2)
    add_body(
        doc,
        "用户不愿意在多个步骤之间来回切换，所以首个版本应把上传、理解、归类和预览放在同一屏完成。上传是入口，自动整理是反馈，时间线可编辑是信任建立的关键。"
    )
    add_heading(doc, "2.4 以“可再生成”作为内容终点", level=2)
    add_body(
        doc,
        "同一份旅行数据应该能输出多种形态：Plog 偏图文排版，Vlog 偏镜头节奏，后续还可以扩展成旅行总结、路线卡、收藏页和年度回顾。内容资产一旦结构化，后续产物就不会彼此割裂。"
    )

    add_heading(doc, "3. Web 上传版交互结构", level=1)
    add_body(
        doc,
        "当前 demo 的三栏结构不是随意排版，而是对应用户最短路径：左边先收素材，中间确认旅行事实，右边把事实变成作品。这样可以把“上传、理解、编辑、输出”放进同一个闭环。"
    )
    add_key_value_table(
        doc,
        [
            ("界面区", "作用", "设计要点"),
            ("左侧素材栏", "完成拖拽或选择上传，展示全部原图与新素材状态。", "先把素材收齐，再谈内容理解；新上传图片要有即时反馈。"),
            ("中央时间线", "展示旅行事件、地点、时间与故事文本。", "以事件为单位编辑，避免用户在照片堆里寻找上下文。"),
            ("右侧生成面板", "生成 Plog 或 Vlog，并提供风格切换。", "输出层和编辑层并存，减少跳转成本。"),
            ("新建旅行弹窗", "创建新的旅行项目。", "先命名，再上传，保证项目边界清楚。"),
        ],
        [1.25, 2.55, 2.70],
    )

    add_heading(doc, "4. 技术方案", level=1)
    add_body(
        doc,
        "技术上建议把系统拆成五层：前端交互、素材接入、内容理解、内容生成、数据存储与任务编排。这样既能支持当前 Web demo，也能平滑过渡到后端上传和更强的 AI 能力。"
    )
    tech_table = doc.add_table(rows=1, cols=4)
    set_table_layout(tech_table, [1.05, 2.15, 1.9, 1.4])
    style_table(tech_table, header_fill="E8EEF5")
    headers = ["层", "职责", "建议实现", "说明"]
    for idx, text in enumerate(headers):
        put_text(tech_table.rows[0].cells[idx], text, bold=True, size=10.2, color="0B2545", align=WD_ALIGN_PARAGRAPH.CENTER)
    tech_rows = [
        ("前端交互层", "负责上传、时间线编辑、风格切换、预览和导出入口。", "React + Vite + 组件化状态管理。", "先把体验做顺，再逐步接后端。"),
        ("素材接入层", "接收图片、生成缩略图、保存原图与元数据。", "Web Upload API / 后端上传接口 + 对象存储。", "上传稳定性决定产品可信度。"),
        ("内容理解层", "识别地点、人物、场景、餐饮、交通等事件线索。", "EXIF、OCR、图片多模态模型、规则与 LLM 协同。", "先规则后模型，避免完全黑盒。"),
        ("内容生成层", "把事件串成旅行故事，并输出文案、Plog、Vlog 脚本。", "模板引擎 + LLM + 可配置风格。", "模板负责稳定，模型负责变化。"),
        ("数据与任务层", "管理旅行、素材、事件、生成任务和版本。", "Postgres / Supabase + 队列任务 + 可追踪状态。", "保证可回溯和可重试。"),
    ]
    for row in tech_rows:
        cells = tech_table.add_row().cells
        for idx, text in enumerate(row):
            put_text(
                cells[idx],
                text,
                bold=(idx == 0),
                size=10.0 if idx != 3 else 9.8,
                color="2B2F33" if idx != 0 else "0B2545",
                align=WD_ALIGN_PARAGRAPH.LEFT if idx != 0 else WD_ALIGN_PARAGRAPH.CENTER,
            )

    add_heading(doc, "5. MVP 分阶段路线", level=1)
    roadmap = doc.add_table(rows=1, cols=3)
    set_table_layout(roadmap, [1.15, 1.9, 3.45])
    style_table(roadmap, header_fill="F2F4F7")
    for idx, text in enumerate(["阶段", "目标", "交付物"]):
        put_text(roadmap.rows[0].cells[idx], text, bold=True, size=10.2, color="0B2545", align=WD_ALIGN_PARAGRAPH.CENTER)
    roadmap_rows = [
        ("阶段 1", "先验证 Web 上传版是否能稳定收图和成线。", "上传、拖拽、时间线编辑、Plog/Vlog 预览、基础导出。"),
        ("阶段 2", "把前端演示变成可用产品。", "后端上传、对象存储、EXIF/OCR、任务队列、用户登录。"),
        ("阶段 3", "把内容生成做成真正的差异化能力。", "多模态理解、自动分段、脚本生成、模板视频渲染、分享页。"),
    ]
    for row in roadmap_rows:
        cells = roadmap.add_row().cells
        for idx, text in enumerate(row):
            put_text(
                cells[idx],
                text,
                bold=(idx == 0),
                size=10.0,
                color="2B2F33" if idx != 0 else "0B2545",
                align=WD_ALIGN_PARAGRAPH.LEFT if idx != 0 else WD_ALIGN_PARAGRAPH.CENTER,
            )

    add_heading(doc, "6. 当前 Web Demo 已实现的内容", level=1)
    add_body(
        doc,
        "当前版本已经完成了一个可演示的前端闭环：支持图片上传与拖拽上传，自动把新素材并入旅行素材栏；支持时间线编辑，让用户直接改写每一段旅行故事；支持 Plog 与 Vlog 两种输出形态的预览与生成；支持新建旅行弹窗和移动端自适应布局。"
    )
    add_body(
        doc,
        "需要明确的是，这一版的“自动理解”仍然是前端模拟逻辑，目的是先把产品交互、信息结构和视觉语言跑通。下一步真正的价值点，不是继续堆更多按钮，而是把上传、理解和生成三段链路接成可追踪的后端流程。"
    )

    add_heading(doc, "7. 我建议的下一步", level=1)
    add_body(
        doc,
        "如果继续推进，我会优先补三件事：第一，上传走后端和对象存储，保证图片可持久化；第二，接入 EXIF、OCR 和基础图像理解，把时间线从“演示”变成“可信”；第三，做一套模板化导出，把 Plog 和 Vlog 变成真正可下载、可分享的成果。"
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
